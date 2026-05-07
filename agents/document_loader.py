"""
Agent: Document Loader
=======================
Parses PDF and Word (.docx) documents, splits into chunks,
and indexes them into ChromaDB + rebuilds BM25 index.

Supported formats:
  - PDF  → pypdf (fast) with docling fallback (scans / complex layouts)
  - DOCX → python-docx

Usage:
  from agents.document_loader import load_document
  result = load_document("docs/report.pdf")
  result = load_document("docs/report.docx")
"""

import re
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Re-use the same vectorstore & BM25 from existing agents
from agents.vector_db  import vectorstore
from agents.bm25_index import get_bm25, _tokenize
import agents.bm25_index as _bm25_module

DOCS_DIR   = Path(__file__).parent.parent / "docs"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 100

splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", ". ", " "],
)


# ══════════════════════════════════════
#  Parsers
# ══════════════════════════════════════

def _parse_pdf(path: Path) -> List[str]:
    """Extract text pages from PDF using pypdf, fall back to docling."""
    pages = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        # If less than 60% of pages extracted, try docling
        if len(pages) < len(reader.pages) * 0.6:
            raise ValueError("pypdf extracted too little text, trying docling")
        print(f"  [Loader] PDF parsed with pypdf: {len(pages)} pages")
        return pages
    except Exception as e:
        print(f"  [Loader] pypdf issue ({e}), trying docling...")

    try:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result    = converter.convert(str(path))
        text      = result.document.export_to_markdown()
        # Split by markdown headings or double newlines as "pages"
        pages = [p.strip() for p in re.split(r"\n{3,}", text) if p.strip()]
        print(f"  [Loader] PDF parsed with docling: {len(pages)} sections")
        return pages
    except Exception as e:
        print(f"  [Loader] docling failed: {e}")
        return []


def _parse_docx(path: Path) -> List[str]:
    """Extract paragraphs from Word document."""
    from docx import Document as DocxDocument
    doc    = DocxDocument(str(path))
    pages  = []
    buffer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if buffer:
                pages.append("\n".join(buffer))
                buffer = []
        else:
            buffer.append(text)

    if buffer:
        pages.append("\n".join(buffer))

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            pages.append("\n".join(rows))

    print(f"  [Loader] DOCX parsed: {len(pages)} sections")
    return pages


# ══════════════════════════════════════
#  Chunking
# ══════════════════════════════════════

def _chunk(pages: List[str], source: str, page_label: str = "стр.") -> List[Document]:
    """Split pages into chunks with source metadata."""
    docs = []
    for page_num, text in enumerate(pages, 1):
        chunks = splitter.split_text(text)
        for chunk in chunks:
            if chunk.strip():
                docs.append(Document(
                    page_content=chunk,
                    metadata={"source": source, "page": page_num, "page_label": page_label},
                ))
    return docs


# ══════════════════════════════════════
#  Index into ChromaDB + rebuild BM25
# ══════════════════════════════════════

def _delete_by_source(source: str):
    """Delete all existing chunks for a given source from ChromaDB."""
    col = vectorstore._collection
    results = col.get(where={"source": source}, include=[])
    ids = results.get("ids", [])
    if ids:
        vectorstore.delete(ids=ids)  # langchain interface persists to disk
        print(f"  [Loader] Deleted {len(ids)} old chunks for '{source}'")


def _index(docs: List[Document]) -> int:
    """Delete old chunks for source, add new ones, rebuild BM25 cache."""
    if not docs:
        return 0
    source = docs[0].metadata.get("source", "")
    if source:
        _delete_by_source(source)
    batch_size = 500
    for i in range(0, len(docs), batch_size):
        batch = docs[i:i + batch_size]
        vectorstore.add_documents(batch)
        print(f"  [Loader] Indexed batch {i//batch_size + 1}: {len(batch)} chunks")

    # Rebuild BM25 cache
    _bm25_module._bm25_cache = None
    _bm25_module._docs_cache = None
    get_bm25()  # triggers rebuild

    total = vectorstore._collection.count()
    print(f"  [Loader] Indexed {len(docs)} chunks. Total in DB: {total}")
    return len(docs)


# ══════════════════════════════════════
#  Public API
# ══════════════════════════════════════

def load_document(filename: str) -> dict:
    """
    Parse a PDF or DOCX file from docs/ folder, chunk it, index it.

    Returns:
        {
          "filename": str,
          "format":   "pdf" | "docx",
          "pages":    int,
          "chunks":   int,
          "total_in_db": int,
          "status":   "ok" | "error",
          "message":  str,
        }
    """
    path = DOCS_DIR / filename
    if not path.exists():
        return {"filename": filename, "status": "error",
                "message": f"File not found: {path}"}

    ext    = path.suffix.lower()
    pages  = []

    if ext == ".pdf":
        pages = _parse_pdf(path)
        fmt   = "pdf"
        page_label = "p."
    elif ext in (".docx", ".doc"):
        pages = _parse_docx(path)
        fmt   = "docx"
        page_label = "sec."
    else:
        return {"filename": filename, "status": "error",
                "message": f"Unsupported format: {ext}. Use .pdf or .docx"}

    if not pages:
        return {"filename": filename, "format": fmt, "status": "error",
                "message": "Could not extract text from the document"}

    docs        = _chunk(pages, source=filename, page_label=page_label)
    indexed     = _index(docs)
    total_in_db = vectorstore._collection.count()

    return {
        "filename":    filename,
        "format":      fmt,
        "pages":       len(pages),
        "chunks":      indexed,
        "total_in_db": total_in_db,
        "status":      "ok",
        "message":     f"Loaded {indexed} chunks from {len(pages)} pages/sections",
    }


def list_docs() -> List[str]:
    """Return all PDF and DOCX files in docs/ folder."""
    return [
        f.name for f in sorted(DOCS_DIR.iterdir())
        if f.suffix.lower() in (".pdf", ".docx", ".doc")
    ]
